from docx import Document
from sys import stdin

a1 = []
a2 = []
for i in range(2):
    a1.append(input())
for el in stdin:
    a2.append(el)
document = Document()
for el in a2:
    document.add_heading("Дорогая " + el + "\n")
    document.add_paragraph("приглашаем тебя на праздновапние 8 марта которое состоится " + a1[0])
    document.add_paragraph(a1[1].lower())
    document.add_page_break()
    document.save('test4.docx')
