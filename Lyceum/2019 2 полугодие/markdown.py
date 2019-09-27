from docx import Document


def markdown_to_docx(text):
    a = text.split('\n')
    document = Document()
    document.add_heading(a[0], 0)
    # for i in range(35, len(a)):
    #     print(a[i])
    for i in range(1, len(a)):
        el = a[i]
        if el[:1] == '#':
            document.add_heading(el.replace(' #', '').replace('# ', '').replace('#', ''), level=el[:7].count('#'))
        elif el.replace('-', '*').replace('+', '*')[:2] == '* ':
            p = document.add_paragraph()
            p.style = 'List Bullet'
            r = p.add_run()
            r.add_text(' ' + el[2:])
        elif el[:1].isdigit():
            p = document.add_paragraph(style='List Number')
            # p.style = 'List Number'
            r = p.add_run()
            r.add_text(' ' + el[3:])
        elif el[:1] == '*':
            p = document.add_paragraph()
            if el.startswith('***'):
                r = p.add_run(el[4:-3])
                r.bold = True
                r.italic = True
            elif el.startswith('**'):
                r = p.add_run(el[3:-2])
                r.bold = True
            elif el.startswith('*'):
                r = p.add_run(el[2:-1])
                r.bold = True
        else:
            p = document.add_paragraph()
            r = p.add_run()
            r.add_text(el)

    document.save('res.docx')


#
# f = open('ttt.txt', 'r')
# markdown_to_docx(f.read())
